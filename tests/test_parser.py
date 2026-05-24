from __future__ import annotations

import io

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.errors import ParseError
from src.parser import (
    ExpenseItem,
    _extract_amount,
    _extract_from_paragraphs,
    _extract_from_tables,
    _normalise_amount,
    _parse_text_line,
    parse_items,
)


def _doc_with_paragraphs(*lines: str) -> Document:
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    return doc


def _doc_with_table(headers: list[str], rows: list[list[str]]) -> Document:
    doc = Document()
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            table.rows[1 + r_idx].cells[c_idx].text = cell_text
    return doc


def _save_and_load(doc: Document) -> Document:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Document(buf)


class TestNormaliseAmount:
    def test_gbp_with_comma(self):
        assert _normalise_amount("£3,500.00") == 3500.0

    def test_usd_no_decimal(self):
        assert _normalise_amount("$1,200") == 1200.0

    def test_eur_with_space(self):
        assert _normalise_amount("€ 99.9") == pytest.approx(99.9)

    def test_bare_integer(self):
        assert _normalise_amount("42") == 42.0

    def test_bare_float(self):
        assert _normalise_amount("1234.56") == pytest.approx(1234.56)

    def test_raises_value_error_on_empty(self):
        with pytest.raises(ValueError):
            _normalise_amount("")

    def test_raises_value_error_on_symbol_only(self):
        with pytest.raises(ValueError):
            _normalise_amount("£")

    def test_raises_value_error_on_text(self):
        with pytest.raises(ValueError):
            _normalise_amount("abc")

    def test_comma_thousands_multi(self):
        assert _normalise_amount("1,000,000.00") == 1000000.0

    def test_gbp_no_decimal(self):
        assert _normalise_amount("£500") == 500.0


class TestParseTextLine:
    def test_amount_after_description(self):
        item = _parse_text_line("Salary £3,500.00")
        assert item is not None
        assert item.raw_text == "Salary"
        assert item.amount == 3500.0

    def test_direction_defaults_to_out(self):
        item = _parse_text_line("Rent 900.00")
        assert item is not None
        assert item.direction == 'out'

    def test_amount_without_currency(self):
        item = _parse_text_line("Rent 900.00")
        assert item is not None
        assert item.raw_text == "Rent"
        assert item.amount == 900.0

    def test_no_amount_returns_none(self):
        assert _parse_text_line("not a valid line") is None

    def test_empty_line_returns_none(self):
        assert _parse_text_line("") is None

    def test_amount_only_returns_none(self):
        assert _parse_text_line("£3,500.00") is None

    def test_description_stripped(self):
        item = _parse_text_line("  Salary   £3,500.00  ")
        assert item is not None
        assert item.raw_text == "Salary"

    def test_returns_expense_item_type(self):
        item = _parse_text_line("Salary £100")
        assert isinstance(item, ExpenseItem)

    def test_multi_word_description(self):
        item = _parse_text_line("Bill - Council Tax £180.00")
        assert item is not None
        assert item.amount == 180.0


class TestExtractFromParagraphs:
    def test_extracts_items_from_plain_paragraphs(self):
        doc = _save_and_load(_doc_with_paragraphs(
            "Salary £3,500.00",
            "Rent 900.00",
            "not a valid line",
        ))
        items = _extract_from_paragraphs(doc)
        assert len(items) == 2

    def test_all_directions_are_out(self):
        doc = _save_and_load(_doc_with_paragraphs("Salary 3500.00", "Rent 900.00"))
        items = _extract_from_paragraphs(doc)
        assert all(i.direction == 'out' for i in items)

    def test_invalid_lines_skipped(self):
        doc = _save_and_load(_doc_with_paragraphs("no number here", "Valid £10.00"))
        items = _extract_from_paragraphs(doc)
        assert len(items) == 1

    def test_empty_doc_returns_empty_list(self):
        doc = _save_and_load(_doc_with_paragraphs())
        items = _extract_from_paragraphs(doc)
        assert items == []


class TestExtractFromTables:
    def test_extracts_money_in_as_direction_in(self):
        doc = _save_and_load(_doc_with_table(
            ["Description", "Money In", "Money Out"],
            [["Salary", "3500.00", ""]],
        ))
        items = _extract_from_tables(doc)
        assert len(items) == 1
        assert items[0].direction == 'in'
        assert items[0].amount == 3500.0
        assert items[0].raw_text == "Salary"

    def test_extracts_money_out_as_direction_out(self):
        doc = _save_and_load(_doc_with_table(
            ["Description", "Money In", "Money Out"],
            [["Rent", "", "900.00"]],
        ))
        items = _extract_from_tables(doc)
        assert len(items) == 1
        assert items[0].direction == 'out'
        assert items[0].amount == 900.0

    def test_mixed_in_and_out_rows(self):
        doc = _save_and_load(_doc_with_table(
            ["Description", "Money In", "Money Out"],
            [
                ["Salary", "3500.00", ""],
                ["Rent", "", "900.00"],
                ["Refund", "50.00", ""],
            ],
        ))
        items = _extract_from_tables(doc)
        assert len(items) == 3
        directions = {i.raw_text: i.direction for i in items}
        assert directions["Salary"] == 'in'
        assert directions["Rent"] == 'out'
        assert directions["Refund"] == 'in'

    def test_table_without_description_column_skipped(self):
        doc = _save_and_load(_doc_with_table(
            ["Date", "Money In", "Money Out"],
            [["01 Jan", "100.00", ""]],
        ))
        items = _extract_from_tables(doc)
        assert items == []

    def test_table_without_amount_column_skipped(self):
        doc = _save_and_load(_doc_with_table(
            ["Description", "Date"],
            [["Salary", "01 Jan"]],
        ))
        items = _extract_from_tables(doc)
        assert items == []

    def test_header_keywords_are_case_insensitive(self):
        doc = _save_and_load(_doc_with_table(
            ["description", "credit", "debit"],
            [["Salary", "3500.00", ""]],
        ))
        items = _extract_from_tables(doc)
        assert len(items) == 1
        assert items[0].direction == 'in'

    def test_empty_description_row_skipped(self):
        doc = _save_and_load(_doc_with_table(
            ["Description", "Money In", "Money Out"],
            [["", "100.00", ""], ["Salary", "3500.00", ""]],
        ))
        items = _extract_from_tables(doc)
        assert len(items) == 1

    def test_row_with_both_amounts_uses_money_in(self):
        doc = _save_and_load(_doc_with_table(
            ["Description", "Money In", "Money Out"],
            [["Ambiguous", "100.00", "50.00"]],
        ))
        items = _extract_from_tables(doc)
        assert len(items) == 1
        assert items[0].direction == 'in'
        assert items[0].amount == 100.0

    def test_amounts_with_currency_symbols(self):
        doc = _save_and_load(_doc_with_table(
            ["Description", "Money In", "Money Out"],
            [["Salary", "£3,500.00", ""]],
        ))
        items = _extract_from_tables(doc)
        assert items[0].amount == 3500.0

    def test_no_tables_returns_empty_list(self):
        doc = _save_and_load(_doc_with_paragraphs("Salary 3500.00"))
        items = _extract_from_tables(doc)
        assert items == []


class TestParseItems:
    def test_extracts_from_table_first(self):
        doc = _save_and_load(_doc_with_table(
            ["Description", "Money In", "Money Out"],
            [["Salary", "3500.00", ""], ["Rent", "", "900.00"]],
        ))
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        # Write to a real temp file so parse_items can open it
        import tempfile, os
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            with open(path, "wb") as f:
                f.write(buf.read())
            items = parse_items(path)
        finally:
            os.unlink(path)

        assert len(items) == 2
        directions = {i.raw_text: i.direction for i in items}
        assert directions["Salary"] == 'in'
        assert directions["Rent"] == 'out'

    def test_falls_back_to_paragraphs_when_no_table(self):
        doc = _save_and_load(_doc_with_paragraphs(
            "Salary 3500.00",
            "Rent 900.00",
        ))
        import tempfile, io, os
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            with open(path, "wb") as f:
                f.write(buf.read())
            items = parse_items(path)
        finally:
            os.unlink(path)

        assert len(items) == 2
        assert all(i.direction == 'out' for i in items)

    def test_raises_parse_error_on_empty_document(self):
        import tempfile, io, os
        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            with open(path, "wb") as f:
                f.write(buf.read())
            with pytest.raises(ParseError):
                parse_items(path)
        finally:
            os.unlink(path)

    def test_raises_parse_error_on_missing_file(self):
        with pytest.raises(ParseError):
            parse_items("/tmp/does_not_exist_xyz.docx")
