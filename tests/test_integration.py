from __future__ import annotations

import csv
import io
import os
import sys
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document as DocxDocument

from src.main import main, run_pipeline


def _make_docx_with_paragraphs(tmp_path, lines, filename="fixture.docx"):
    """Create a .docx file with one paragraph per line; return path string."""
    doc = DocxDocument()
    for line in lines:
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    path = tmp_path / filename
    path.write_bytes(buf.getvalue())
    return str(path)


def _make_docx_with_table(tmp_path, headers, rows, filename="fixture.docx"):
    """Create a .docx file with a single table; return path string."""
    doc = DocxDocument()
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            table.rows[1 + r_idx].cells[c_idx].text = cell_text
    buf = io.BytesIO()
    doc.save(buf)
    path = tmp_path / filename
    path.write_bytes(buf.getvalue())
    return str(path)


@pytest.fixture
def fixture_docx_paragraphs(tmp_path):
    return _make_docx_with_paragraphs(tmp_path, [
        "Salary 3500.00",
        "Rent 900.00",
        "Food Supplies 150.00",
    ])


@pytest.fixture
def fixture_docx_table(tmp_path):
    return _make_docx_with_table(
        tmp_path,
        headers=["Description", "Money In", "Money Out"],
        rows=[
            ["Salary", "3500.00", ""],
            ["Rent", "", "900.00"],
            ["Food Supplies", "", "150.00"],
        ],
    )


class TestIntegrationRunPipeline:
    def test_output_csv_is_created(self, fixture_docx_paragraphs, tmp_path):
        output_csv = str(tmp_path / "output.csv")
        with patch("src.main.convert_pdf", return_value=fixture_docx_paragraphs):
            run_pipeline("dummy.pdf", output_csv)
        assert os.path.exists(output_csv)

    def test_csv_has_header_row(self, fixture_docx_paragraphs, tmp_path):
        output_csv = str(tmp_path / "output.csv")
        with patch("src.main.convert_pdf", return_value=fixture_docx_paragraphs):
            run_pipeline("dummy.pdf", output_csv)
        with open(output_csv, newline="", encoding="utf-8") as f:
            header = next(csv.reader(f))
        assert header == ["section", "category", "total_amount"]

    def test_csv_contains_total_income_row(self, fixture_docx_paragraphs, tmp_path):
        output_csv = str(tmp_path / "output.csv")
        with patch("src.main.convert_pdf", return_value=fixture_docx_paragraphs):
            run_pipeline("dummy.pdf", output_csv)
        with open(output_csv, newline="", encoding="utf-8") as f:
            rows = list(csv.reader(f))
        assert any("Total Income" in row for row in rows)

    def test_csv_contains_total_expenditure_row(self, fixture_docx_paragraphs, tmp_path):
        output_csv = str(tmp_path / "output.csv")
        with patch("src.main.convert_pdf", return_value=fixture_docx_paragraphs):
            run_pipeline("dummy.pdf", output_csv)
        with open(output_csv, newline="", encoding="utf-8") as f:
            rows = list(csv.reader(f))
        assert any("Total Expenditure" in row for row in rows)

    def test_salary_amount_reflected_in_csv(self, fixture_docx_paragraphs, tmp_path):
        output_csv = str(tmp_path / "output.csv")
        with patch("src.main.convert_pdf", return_value=fixture_docx_paragraphs):
            run_pipeline("dummy.pdf", output_csv)
        content = Path(output_csv).read_text(encoding="utf-8")
        assert "3500.00" in content

    def test_no_temp_files_remain_after_pipeline(self, fixture_docx_paragraphs, tmp_path):
        output_csv = str(tmp_path / "output.csv")
        import tempfile
        before = set(Path(tempfile.gettempdir()).glob("*.docx"))
        with patch("src.main.convert_pdf", return_value=fixture_docx_paragraphs):
            run_pipeline("dummy.pdf", output_csv)
        after = set(Path(tempfile.gettempdir()).glob("*.docx"))
        assert after == before

    def test_table_extraction_sets_direction(self, fixture_docx_table, tmp_path):
        output_csv = str(tmp_path / "output.csv")
        with patch("src.main.convert_pdf", return_value=fixture_docx_table):
            unmatched = run_pipeline("dummy.pdf", output_csv)
        assert unmatched == []
        content = Path(output_csv).read_text(encoding="utf-8")
        assert "3500.00" in content

    def test_unmatched_items_in_uncategorised_section(self, tmp_path):
        docx_path = _make_docx_with_paragraphs(tmp_path, [
            "Salary 3500.00",
            "mystery item xyzzy 42.00",
        ])
        output_csv = str(tmp_path / "output.csv")
        with patch("src.main.convert_pdf", return_value=docx_path):
            unmatched = run_pipeline("dummy.pdf", output_csv)
        assert len(unmatched) == 1
        content = Path(output_csv).read_text(encoding="utf-8")
        assert "Uncategorised" in content


class TestIntegrationMainExitBehaviour:
    def test_main_exits_zero_on_valid_input(self, fixture_docx_paragraphs, tmp_path):
        output_csv = str(tmp_path / "output.csv")
        pdf_path = str(tmp_path / "input.pdf")
        Path(pdf_path).write_bytes(b"%PDF-1.4")
        with patch("sys.argv", ["expense-summary", pdf_path, output_csv]):
            with patch("src.main.convert_pdf", return_value=fixture_docx_paragraphs):
                with pytest.raises(SystemExit) as exc_info:
                    main()
        assert exc_info.value.code == 0

    def test_main_creates_csv_on_valid_input(self, fixture_docx_paragraphs, tmp_path):
        output_csv = str(tmp_path / "output.csv")
        pdf_path = str(tmp_path / "input.pdf")
        Path(pdf_path).write_bytes(b"%PDF-1.4")
        with patch("sys.argv", ["expense-summary", pdf_path, output_csv]):
            with patch("src.main.convert_pdf", return_value=fixture_docx_paragraphs):
                with pytest.raises(SystemExit):
                    main()
        assert os.path.exists(output_csv)
