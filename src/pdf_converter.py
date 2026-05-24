from __future__ import annotations

import os
import tempfile

from src.errors import ConversionError


def _convert_with_pdf2docx(pdf_path: str, docx_path: str) -> bool:
    """Run pdf2docx conversion from pdf_path to docx_path. Return True on success."""
    try:
        from pdf2docx import Converter
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        return True
    except Exception:
        return False


def _docx_has_content(docx_path: str) -> bool:
    """Return True if the .docx has at least one table row or non-empty paragraph."""
    try:
        from docx import Document
        doc = Document(docx_path)
        if any(
            any(cell.text.strip() for cell in row.cells)
            for table in doc.tables
            for row in table.rows
        ):
            return True
        return any(para.text.strip() for para in doc.paragraphs)
    except Exception:
        return False


def convert_pdf(pdf_path: str) -> str:
    """Convert PDF to a temp .docx file; return the temp file path.
    Raises ConversionError if conversion fails or yields an empty document."""
    fd, docx_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        if not _convert_with_pdf2docx(pdf_path, docx_path):
            raise ConversionError(f"pdf2docx could not convert: {pdf_path}")
        if not _docx_has_content(docx_path):
            raise ConversionError(f"Conversion produced an empty document: {pdf_path}")
    except ConversionError:
        if os.path.exists(docx_path):
            os.unlink(docx_path)
        raise
    return docx_path
