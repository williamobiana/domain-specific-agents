from __future__ import annotations

import re
from dataclasses import dataclass, field

from src.errors import ParseError


@dataclass
class ExpenseItem:
    raw_text: str
    amount: float
    direction: str = field(default='out')  # 'in' (Money In) or 'out' (Money Out)


_AMOUNT_RE = re.compile(r'[£$€]?\s*[\d,]+(?:\.\d{1,2})?')

# Keywords that indicate a Money In header column
_IN_HEADERS = {'money in', 'credit', 'credits', 'in'}
# Keywords that indicate a Money Out header column
_OUT_HEADERS = {'money out', 'debit', 'debits', 'out'}
# Keywords that indicate the description column
_DESC_HEADERS = {'description', 'details', 'narrative', 'particulars', 'payee'}


def _normalise_amount(raw: str) -> float:
    """Strip currency symbols, commas, and whitespace then convert to float."""
    cleaned = re.sub(r'[£$€,\s]', '', raw)
    return float(cleaned)


def _parse_text_line(line: str) -> ExpenseItem | None:
    """Extract a description and amount from a plain text line; direction defaults to 'out'."""
    match = _AMOUNT_RE.search(line)
    if not match:
        return None
    raw_amount = match.group()
    description = (line[: match.start()] + line[match.end():]).strip()
    if not description:
        return None
    try:
        amount = _normalise_amount(raw_amount)
    except ValueError:
        return None
    return ExpenseItem(raw_text=description, amount=amount, direction='out')


def _find_col(headers: list[str], keywords: set[str]) -> int | None:
    """Return the index of the first header cell whose normalised text is in keywords."""
    for i, h in enumerate(headers):
        norm = h.lower().strip()
        if norm in keywords:
            return i
    return None


def _extract_amount(cell_text: str) -> float | None:
    """Extract a numeric amount from a table cell; return None if empty or unparseable."""
    text = cell_text.strip()
    if not text:
        return None
    match = _AMOUNT_RE.search(text)
    if not match:
        return None
    try:
        return _normalise_amount(match.group())
    except ValueError:
        return None


def _extract_from_tables(doc) -> list[ExpenseItem]:  # type: ignore[type-arg]
    """Extract ExpenseItems from Word tables with Money In / Money Out columns."""
    items: list[ExpenseItem] = []
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [cell.text for cell in table.rows[0].cells]
        desc_col = _find_col(headers, _DESC_HEADERS)
        in_col = _find_col(headers, _IN_HEADERS)
        out_col = _find_col(headers, _OUT_HEADERS)

        # Need at least a description column and one amount column
        if desc_col is None or (in_col is None and out_col is None):
            continue

        for row in table.rows[1:]:
            cells = [cell.text for cell in row.cells]
            if len(cells) <= desc_col:
                continue
            description = cells[desc_col].strip()
            if not description:
                continue

            in_amount = _extract_amount(cells[in_col]) if in_col is not None and in_col < len(cells) else None
            out_amount = _extract_amount(cells[out_col]) if out_col is not None and out_col < len(cells) else None

            if in_amount is not None:
                items.append(ExpenseItem(raw_text=description, amount=in_amount, direction='in'))
            elif out_amount is not None:
                items.append(ExpenseItem(raw_text=description, amount=out_amount, direction='out'))

    return items


def _extract_from_paragraphs(doc) -> list[ExpenseItem]:  # type: ignore[type-arg]
    """Fallback: extract ExpenseItems from plain text paragraphs."""
    items: list[ExpenseItem] = []
    for para in doc.paragraphs:
        item = _parse_text_line(para.text)
        if item is not None:
            items.append(item)
    return items


def parse_items(docx_path: str) -> list[ExpenseItem]:
    """Parse a .docx file into a list of ExpenseItem objects.
    Tries table extraction first, then paragraph extraction.
    Raises ParseError if the result is empty."""
    from docx import Document  # imported here to keep top-level imports clean
    try:
        doc = Document(docx_path)
    except Exception as exc:
        raise ParseError(f"Could not open document: {exc}") from exc

    items = _extract_from_tables(doc)
    if not items:
        items = _extract_from_paragraphs(doc)

    if not items:
        raise ParseError("No expense items found in the document.")
    return items
